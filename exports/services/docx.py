"""Serviços utilitários para gerar DOCX."""

from __future__ import annotations

from html import unescape
from html.parser import HTMLParser
from io import BytesIO
from typing import Any, Dict, List, Tuple

try:
    from docx import Document
except ImportError:  # pragma: no cover - dependencia externa
    Document = None  # type: ignore


class _ParagraphCollector(HTMLParser):
    """Extrai pares (tipo, texto) preservando limites de parágrafos básicos."""

    _BLOCK_TAGS = {
        "p",
        "div",
        "section",
        "article",
        "h1",
        "h2",
        "h3",
        "h4",
        "h5",
        "h6",
    }
    _LIST_TAGS = {"li"}

    def __init__(self) -> None:
        super().__init__()
        self._segments: List[Tuple[str, str]] = []
        self._current_parts: List[str] = []
        self._current_kind = "block"

    @property
    def segments(self) -> List[Tuple[str, str]]:
        self._flush_current()
        return self._segments

    def handle_starttag(self, tag: str, attrs) -> None:  # type: ignore[override]
        tag = tag.lower()
        if tag == "br":
            self._flush_current()
        elif tag in self._LIST_TAGS:
            self._flush_current()
            self._current_kind = "list"
        elif tag in self._BLOCK_TAGS:
            self._flush_current()
            self._current_kind = "block"

    def handle_endtag(self, tag: str) -> None:  # type: ignore[override]
        tag = tag.lower()
        if tag in self._LIST_TAGS:
            self._flush_current()
            self._current_kind = "block"
        elif tag in self._BLOCK_TAGS:
            self._flush_current()

    def handle_data(self, data: str) -> None:  # type: ignore[override]
        text = data.replace("\xa0", " ")
        if not text.strip():
            return
        text = text.strip()
        if self._current_parts:
            last = self._current_parts[-1]
            if last and not last[-1].isspace() and text and text[0] not in ",.;:!?)]}":
                self._current_parts.append(" ")
        self._current_parts.append(text)

    def handle_entityref(self, name: str) -> None:  # type: ignore[override]
        self.handle_data(unescape(f"&{name};"))

    def handle_charref(self, name: str) -> None:  # type: ignore[override]
        self.handle_data(unescape(f"&#{name};"))

    def handle_startendtag(self, tag: str, attrs) -> None:  # type: ignore[override]
        if tag.lower() == "br":
            self._flush_current()

    def _flush_current(self) -> None:
        if not self._current_parts:
            return
        texto = "".join(self._current_parts).strip()
        if texto:
            self._segments.append((self._current_kind, texto))
        self._current_parts = []


def _extract_segments(html: str) -> List[Tuple[str, str]]:
    parser = _ParagraphCollector()
    parser.feed(html)
    parser.close()
    return [(kind, texto) for kind, texto in parser.segments if texto]


def html_to_docx_bytes(context: Dict[str, Any]) -> bytes:
    if Document is None:
        raise RuntimeError("python-docx não está disponível no ambiente")

    document = Document()
    titulo = context.get("titulo") or "Exportação"
    document.add_heading(titulo, level=1)

    conteudo = context.get("conteudo_html", "")
    for kind, texto in _extract_segments(str(conteudo)):
        if not texto:
            continue
        if kind == "list":
            paragrafo = document.add_paragraph(texto)
            try:
                paragrafo.style = "List Bullet"
            except (KeyError, AttributeError):  # pragma: no cover - estilo ausente
                pass
        else:
            document.add_paragraph(texto)

    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer.read()
